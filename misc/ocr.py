import os
import re
import sys

sys.path.append(os.path.join(os.path.dirname(os.path.abspath(__file__)), ".."))

from s3 import get_file_local
from sqlalchemy.orm import Session

from models import SessionLocal
from models.file import File
from models.task import Task
from misc.const import PIC_EXT, WORD_EXT


def get_text_pdf(file_path: str) -> str:
    import fitz

    document = fitz.Document(file_path)
    texts = []

    for pg in range(document.page_count):
        page = document[pg]
        text = page.get_text("text")

        texts.append(text)

    if len((text := "\n".join(texts))) > 64:
        return text

    texts.clear()

    import cv2
    import numpy as np
    from paddleocr import PaddleOCR
    from PIL import Image

    ocr = PaddleOCR(use_angle_cls=True, lang="ch")
    for pg in range(document.page_count):
        UpdateStatus(pg / document.page_count)
        page = document[pg]

        mat = fitz.Matrix(2, 2)
        pm = page.get_pixmap(matrix=mat, alpha=False)
        # if width or height > 2000 pixels, don't enlarge the image
        if pm.width > 2000 or pm.height > 2000:
            pm = page.get_pixmap(matrix=fitz.Matrix(1, 1), alpha=False)

        img = Image.frombytes("RGB", [pm.width, pm.height], pm.samples)
        img = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)

        result = ocr.ocr(img, cls=True)
        fulltext = []
        for idx in range(len(result)):
            res = result[idx]
            for line in res:
                fulltext.append(line[1][0])
        texts.append("\n".join(fulltext))

    return "\n".join(texts)


def get_text_image(file_path: str) -> str:
    from paddleocr import PaddleOCR

    ocr = PaddleOCR(use_angle_cls=True, lang="ch")
    result = ocr.ocr(file_path, cls=True)
    fulltext = []
    for idx in range(len(result)):
        res = result[idx]
        for line in res:
            fulltext.append(line[1][0])

    return "\n".join(fulltext)


def get_text_docx(file_path: str) -> str:
    import docx

    doc = docx.Document(file_path)
    fulltext = []
    for para in doc.paragraphs:
        fulltext.append(para.text)
    return "\n".join(fulltext)


def get_text(file_path: str, ext: str) -> str:
    if ext == "pdf":
        text = get_text_pdf(file_path)
    elif ext in WORD_EXT:
        text = get_text_docx(file_path)
    elif ext in PIC_EXT:
        text = get_text_image(file_path)
    else:
        text = ""

    text = re.sub("([^\u4e00-\u9fa5\u0041-\u005a\u0061-\u007a\u0030-\u0039])", "", text)

    return text


def WriteOCR(fid: str):
    global db

    file = db.query(File).filter_by(fid=fid).first()

    file_path = get_file_local(file.ext, file.md5)
    text = get_text(file_path, file.ext)

    db.query(File).filter_by(fid=fid).update({"ocr": text})

    os.remove(file_path)
    db.commit()


def UpdateStatus(status: float):
    global task, db

    task.status = status
    db.commit()


if __name__ == "__main__":
    db: Session = SessionLocal()

    task = db.query(Task).filter_by(type="ocr").order_by(Task.created.asc()).first()
    while task:
        try:
            WriteOCR(task.data)
        except:
            pass
        db.delete(task)
        db.commit()
        task = db.query(Task).filter_by(type="ocr").order_by(Task.created.asc()).first()
